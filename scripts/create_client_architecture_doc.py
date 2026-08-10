"""Create a client-ready architecture brief for the MyCodeAgent platform."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "MyCodeAgent_High_Level_Architecture.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
INK = "1F2937"
GREEN = "1F5D42"
GOLD = "7A5A00"


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D8DEE8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def write_cell(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        write_cell(cell, header, bold=True, color=NAVY, size=9.5)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            write_cell(cell, value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text, emphasis=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if not emphasis:
        run = paragraph.add_run(text)
        set_font(run, size=10.8, color=INK)
        return paragraph
    before, highlighted, after = emphasis
    for value, bold, color in ((before, False, INK), (highlighted, True, NAVY), (after, False, INK)):
        run = paragraph.add_run(value)
        set_font(run, size=10.8, color=color, bold=bold)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.08
        run = paragraph.add_run(item)
        set_font(run, size=10.5, color=INK)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.08
        run = paragraph.add_run(item)
        set_font(run, size=10.5, color=INK)


def add_callout(doc, label, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="B7C9DE")
    cell = table.cell(0, 0)
    shade(cell, "F4F7FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(f"{label}  ")
    set_font(run, size=10.3, color=color, bold=True)
    run = p.add_run(text)
    set_font(run, size=10.3, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_flow(doc):
    stages = [
        ("1", "Task intake", "A ready task in TODO.md defines scope, workspace, tests, and delivery rules."),
        ("2", "Isolated execution", "Python creates an optional task branch/worktree and prepares the bounded work order."),
        ("3", "Implementation", "The selected implementation agent writes code and tests within the task workspace."),
        ("4", "Verification", "Tests run and the reviewer agent assesses source, tests, and task acceptance criteria."),
        ("5", "Remediation", "Only review findings are returned to the implementer for a bounded corrective pass."),
        ("6", "Delivery", "Python creates changelog, commit, push, and PR only after final APPROVED status."),
    ]
    for index, (number, title, detail) in enumerate(stages):
        table = doc.add_table(rows=1, cols=2)
        set_table_geometry(table, [720, 8640])
        set_table_borders(table, color="C7D4E2")
        left, right = table.rows[0].cells
        shade(left, BLUE if index < 5 else GREEN)
        write_cell(left, number, bold=True, color=WHITE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(right, "F7F9FC")
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title + " — ")
        set_font(run, size=10.2, color=NAVY, bold=True)
        run = p.add_run(detail)
        set_font(run, size=10.2, color=INK)
        if index < len(stages) - 1:
            arrow = doc.add_paragraph()
            arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow.paragraph_format.space_before = Pt(0)
            arrow.paragraph_format.space_after = Pt(0)
            run = arrow.add_run("↓")
            set_font(run, size=12, color=MUTED, bold=True)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("MYCODEAGENT | HIGH-LEVEL ARCHITECTURE")
    set_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Client architecture overview | Confidential working draft")
    set_font(run, size=8, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("MyCodeAgent")
    set_font(run, size=27, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("High-Level Architecture for Governed Multi-Agent Software Delivery")
    set_font(run, size=15, color=MUTED)
    add_callout(
        doc,
        "Purpose",
        "Provide a controlled, auditable path from a clearly defined software task to reviewed code and a pull request—while preserving human governance over scope, policy, and delivery.",
    )
    add_table(
        doc,
        ["Audience", "Document purpose", "Scope"],
        [["Business and technology stakeholders", "Explain the operating model and control points", "High-level architecture; not a product specification or security certification"]],
        [2100, 3400, 3860],
    )

    add_heading(doc, "1. Problem statement")
    add_body(doc, "Engineering teams increasingly use AI assistants to accelerate routine software delivery, but unstructured use introduces practical risks: inconsistent task interpretation, missing tests, uncontrolled tool actions, long-running loops, weak review evidence, and unclear accountability for changes sent to source control.")
    add_body(doc, "The challenge is not only generating code. The platform must coordinate specialized agents through a repeatable delivery process, enforce boundaries before side effects occur, and create evidence that a client can inspect after each task.")
    add_table(
        doc,
        ["Business need", "Architecture response"],
        [
            ["Repeatable execution", "A task contract in TODO.md supplies the workspace, acceptance criteria, tests, and lifecycle status."],
            ["Separation of duties", "Implementation and review can be delegated to separate model providers and agents."],
            ["Controlled delivery", "Python—not an LLM—performs deterministic changelog, Git, push, and pull-request operations."],
            ["Operational visibility", "Each task produces a time-stamped trace with stage outcomes and next actions."],
        ],
        [2900, 6460],
    )

    add_heading(doc, "2. Solution architecture")
    add_body(doc, "MyCodeAgent is a governed agent-delivery harness. Omnigent coordinates the model-driven stages; a Python control plane owns task parsing, process limits, status transitions, logs, worktree isolation, and final delivery. This separation makes the LLM valuable for judgment-intensive code and review work while retaining deterministic controls around side effects.")
    add_table(
        doc,
        ["Layer", "Primary components", "Responsibility"],
        [
            ["Task governance", "TODO.md; task schema; status rules", "Defines what can be worked on and what completion means."],
            ["Agent orchestration", "coding_agent.yaml; Omnigent", "Sequences implementation, testing, review, and bounded remediation."],
            ["Model execution", "Claude implementation agent; Codex review agent (configurable)", "Performs code changes and independent quality assessment."],
            ["Control plane", "Python CLI and workflow modules", "Launches stages, enforces limits, records traces, and validates reports."],
            ["Delivery automation", "Git helper; GitHub CLI/API", "Creates a PR only after the approved review gate is met."],
        ],
        [1800, 2950, 4610],
    )
    add_callout(doc, "Design principle", "LLMs propose, implement, test, and review; deterministic Python services validate state and perform external side effects.")

    add_heading(doc, "3. Omnigent harness and multi-agent operating model")
    add_body(doc, "Omnigent acts as the agent harness: it interprets the workflow configuration, invokes the defined agent capabilities, and carries the task context between stages. The harness does not receive unrestricted authority to alter GitHub or task policy; those operations remain in the Python delivery layer.")
    add_table(
        doc,
        ["Role", "Illustrative agent", "Accountability"],
        [
            ["Implementation agent", "Claude", "Reads the bounded task context; creates or changes code and tests; reports files changed and validation results."],
            ["Review agent", "Codex", "Independently reads the actual workspace and tests; checks task acceptance criteria, guardrails, and regressions; returns APPROVED or actionable findings."],
            ["Orchestrator", "Omnigent", "Runs the configured sequence and passes complete review findings to a single targeted remediation pass when allowed."],
            ["Control plane", "Python", "Applies timeouts, process cleanup, status updates, trace writing, Git policy checks, and post-approval delivery."],
        ],
        [1900, 1700, 5760],
    )
    add_bullets(doc, [
        "Provider choice is configurable: Claude can be assigned implementation while Codex performs review, or the assignments can be changed to suit cost, capability, and governance requirements.",
        "The reviewer is intentionally independent from implementation. It evaluates source and tests directly rather than relying only on a generated summary or Git diff.",
        "A review finding is fed back verbatim to the implementer for a scoped correction, then verified through a final review. The workflow is bounded to avoid uncontrolled loops.",
    ])

    add_heading(doc, "4. Policies and guardrails")
    add_body(doc, "Controls are applied at two levels: platform-level agent policy and task-level delivery policy. Together they constrain both how agents act and what an individual task is permitted to change.")
    add_table(
        doc,
        ["Control level", "Examples", "Why it matters"],
        [
            ["Agent-level policy", "Serial stage ordering; one tool action at a time; explicit terminal report; bounded remediation; no agent-owned Git/PR execution.", "Prevents ambiguous hand-offs, parallel side effects, and unbounded agent loops."],
            ["Task-level policy", "Declared source workspace; acceptance criteria; named test command; status lifecycle; task-specific safety and runtime constraints.", "Keeps agents within the intended business and technical scope."],
            ["Runtime guardrails", "Timeouts; idle-watchdog handling; subprocess/process-group cleanup; output/report validation; redacted task logs.", "Makes long-running execution observable and recoverable."],
            ["Delivery guardrails", "Final APPROVED gate; approved remote and identity; explicit changelog/commit/push/PR sequence; worktree branch isolation.", "Limits external changes to validated, auditable conditions."],
        ],
        [1800, 4400, 3160],
    )
    add_callout(doc, "Important boundary", "Guardrails reduce operational risk but do not replace secure coding review, secrets management, dependency scanning, or human approval for high-impact releases.", color=GOLD)

    add_heading(doc, "5. End-to-end task process")
    add_body(doc, "The task lifecycle is designed as a gated sequence. The Python control plane owns status changes and delivery; Omnigent owns the LLM collaboration inside the approved task boundary.")
    add_flow(doc)
    add_table(
        doc,
        ["Status", "Meaning", "Typical exit"],
        [
            ["ready", "Task is eligible for selection.", "Submission starts."],
            ["working", "Implementation/review stages are active.", "Implementation, review, or failure result."],
            ["implemented", "Code and tests were produced; review may still be pending.", "Review is started."],
            ["reviewed", "Review completed but delivery is not approved or not requested.", "Remediate, manually inspect, or deliver after approval."],
            ["delivered", "Approved change was committed, pushed, and a PR was created.", "Normal terminal state."],
            ["failed", "A stage, timeout, or validation gate failed.", "Investigate trace, correct configuration, then retry."],
        ],
        [1600, 4950, 2810],
    )

    add_heading(doc, "6. Implementation detail at a high level")
    add_table(
        doc,
        ["Component", "High-level responsibility"],
        [
            ["CLI and submission service", "Exposes operating modes: implement only; implement plus review; or implement, review, and deliver after approval."],
            ["Task service", "Parses TODO.md into validated task records; confirms the workspace and test command before launch."],
            ["Runner and tracing", "Starts the Omnigent workflow, streams controlled output, captures task-specific traces, and enforces timeout/cleanup behavior."],
            ["Worktree service", "Creates an isolated task branch/worktree from the configured base branch, allowing a task PR to remain independent of other active work."],
            ["Delivery service", "Checks final review status and invokes a deterministic helper to update changelog, commit, push, and open a PR with a task-derived title."],
        ],
        [2700, 6660],
    )

    add_heading(doc, "7. Current scope and next goal: queued execution")
    add_body(doc, "The current recommended operating model processes one ready task at a time, especially for mode 3 (implementation, review, and PR). This provides the clearest audit trail while the review-harness and worktree controls are stabilized.")
    add_body(doc, "The next planned capability is a managed queue that selects multiple ready tasks and processes them sequentially overnight. Each task would receive its own worktree, trace, time budget, final status, and delivery decision before the next task begins.")
    add_table(
        doc,
        ["Queue objective", "Proposed approach", "Required safeguards"],
        [
            ["Process a backlog without manual re-launch", "Read ready tasks in priority/order; execute one task to terminal status before selecting the next.", "Per-task timeout, isolated worktree, durable trace, failure continuation policy, and no shared uncommitted state."],
            ["Support several PRs safely", "Create a distinct branch and worktree for each task from the selected base branch.", "Branch collision checks, remote/identity approval, PR status verification, and cleanup/retention rules."],
            ["Control model cost", "Set a per-stage and per-task time/token budget; stop after the bounded remediation cycle.", "Rate limits, concurrency limit of one initially, clear retry policy, and operator-visible queue ledger."],
        ],
        [2500, 3600, 3260],
    )
    add_callout(doc, "Recommended next milestone", "Implement a sequential queue first. Add parallel task execution only after isolation, provider rate limits, trace completeness, and delivery governance have been demonstrated under load.", color=GREEN)

    add_heading(doc, "8. Client decision points")
    add_numbered(doc, [
        "Select the implementation and review model providers, including permitted repositories and data-handling constraints.",
        "Confirm task policy: required acceptance criteria, test standards, approval thresholds, and maximum execution/time budgets.",
        "Define delivery governance: who may authorize PR creation, which remotes are permitted, and whether human approval is needed before merge.",
        "Agree on observability and retention requirements for task traces, reports, and worktree artifacts.",
        "Approve the phased roadmap: single-task controlled delivery, sequential queue, then optional parallel execution.",
    ])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
